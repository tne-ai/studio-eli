from tne.TNE import TNE
import markdown
from docx import Document
from docx.shared import Pt

session = TNE(uid=UID, bucket_name=BUCKET, project=PROJECT, version=VERSION)
text = PROCESS_INPUT
# Load markdown file
# md_file = session.get_object("inventory-reco.md")
md_file = """

# Inventory Risk Summary

| Metric              | Value                               |
|:--------------------|:------------------------------------|
| Product Category    | Crossbody bag                       |
| At-Risk Styles      | Top 10 high-stock, low-sales styles |
| Total Value at Risk | 199,365.56                          |

I am recommending that we reduce inventory of Handbag Crossbody bag Style 50012.

Here are the justifications:
- The revenue for this bag is underperforming the sales of the other Crossbody bags. See the table showing that Handbag Crossbody bag Style 50012 is the lowest performer:

I'll create a more focused table with just those columns.
# Inventory Risk Summary

| Metric              | Value                               |
|:--------------------|:------------------------------------|
| Product Category    | Crossbody bag                       |
| At-Risk Styles      | Top 10 high-stock, low-sales styles |
| Total Value at Risk | 199,365.56                          |

I am recommending that we reduce inventory of Handbag Crossbody bag Style 50012.

Here are the justifications:
- The revenue for this bag is underperforming the sales of the other Crossbody bags. See the table showing that Handbag Crossbody bag Style 50012 is the lowest performer:

| Style Colour Desc | Product Launch Date | Distinct Store Count | Max Weeks Selling | Total Sales | Rate Of Sale | Average Weekly Sales | End Inv | Sell Through | Total Cost |
|-------------------|-------------------:|--------------------:|------------------:|------------:|-------------:|-------------------:|----------:|-------------:|------------:|
| Handbag Crossbody bag Style 50012 | 21/08/2019 | 281 | 44 | 378 | 0.03 | 8 | 280 | 57 | 12,448.80 |
| Handbag Crossbody bag Style 50233 | 7/1/22 | 279 | 44 | 390 | 0.03 | 8 | 254 | 61 | 21,201.38 |
| Handbag Crossbody bag Style 46672 | 12/1/21 | 270 | 44 | 761 | 0.06 | 17 | 474 | 62 | 30,615.66 |
| Handbag Crossbody bag Style 40761 | 19/04/2020 | 272 | 44 | 188 | 0.02 | 4 | 95 | 66 | 4,389.95 |
| Handbag Crossbody bag Style 46724 | - | 284 | 44 | 651 | 0.05 | 14 | 244 | 73 | 22,472.40 |
| Handbag Crossbody bag Style 40077 | 4/6/12 | 281 | 44 | 810 | 0.07 | 18 | 435 | 65 | 9,269.85 |
| Handbag Crossbody bag Style 46648 | 8/2/22 | 271 | 44 | 899 | 0.08 | 20 | 536 | 63 | 73,153.28 |
| Handbag Crossbody bag Style 36694 | 6/9/22 | 263 | 44 | 435 | 0.04 | 9 | 166 | 72 | 11,108.72 |
| Handbag Crossbody bag Style 40758 | 21/05/2018 | 272 | 44 | 295 | 0.02 | 6 | 78 | 79 | 5,455.32 |
| Handbag Crossbody bag Style 51801 | 3/12/20 | 272 | 9 | 159 | 0.06 | 17 | 270 | 37 | 9,250.20 |
| **Totals** | | | | 4,966 | | | 2,832 | | 199,365.56 |
- Our inventory of Handbag Crossbody bag Style 50012 is 280.
- This bag has been stocked since 21/08/2019 and we have sold only 378 units.

Our recommendation is to put this product on clearance.

Note: Products are selected for the clearance recommendation based on a combined ranking of worst performance, where "worst" is defined as the highest sum of the individual ranks for low sales rate (units sold / weeks available) and high ending inventory. The top 10 worst-performing products are shown in the table, sorted by this combined rank, with the absolute worst performer appearing first.
- Our inventory of Handbag Crossbody bag Style 50012 is 280.
- This bag has been stocked since 21/08/2019 and we have sold only 378 units.

Our recommendation is to put this product on clearance.

Note: Products are selected for the clearance recommendation based on a combined ranking of worst performance, where "worst" is defined as the highest sum of the individual ranks for low sales rate (units sold / weeks available) and high ending inventory. The top 10 worst-performing products are shown in the table, sorted by this combined rank, with the absolute worst performer appearing first.
"""

# Convert markdown to HTML
html = markdown.markdown(md_file)

# Create a new Word document
document = Document()

# Add a title
title = "Inventory Risk Summary"
document.add_heading(title, 0)

# Add the markdown content
for line in html.splitlines():
    if line.startswith("<h1>"):
        document.add_heading(line[4:-5], 1)
    elif line.startswith("<h2>"):
        document.add_heading(line[5:-6], 2)
    elif line.startswith("<p>"):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line[3:-4])
        run.font.size = Pt(12)
    elif line.startswith("<table>"):
        table = document.add_table(rows=1, cols=10, style='Table Grid')
        for row in html.splitlines()[html.splitlines().index(line)+1:]:
            if row.startswith("</table>"):
                break
            row_cells = table.add_row().cells
            for i, cell in enumerate(row.split("|")[1:-1]):
                row_cells[i].text = cell.strip()

# Save the document to a file
document.save("inventory-reco.docx")

# Open the file in binary mode and upload it
with open("inventory-reco.docx", "rb") as f:
    session.upload_object("inventory-reco.docx", f.read())

result = text
